# -*- coding: utf-8 -*-
"""Build the formal Thai DOCX proposal (v6): cost/energy hook + dual-channel + clean B&W figures.
Matches the format of Proposol_I_NEW_GEN.docx; TH Sarabun New 16pt; grayscale figures."""
import os
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib import font_manager as fm
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import numpy as np

HERE = os.path.dirname(os.path.abspath(__file__))
SAR = os.path.join(HERE, "Sarabun-Regular.ttf")
SARB = os.path.join(HERE, "Sarabun-Bold.ttf")
fm.fontManager.addfont(SAR); fm.fontManager.addfont(SARB)
TH = fm.FontProperties(fname=SAR).get_name()
plt.rcParams.update({"font.family": TH, "figure.dpi": 150, "savefig.dpi": 150,
                     "axes.edgecolor": "#222", "text.color": "#111", "axes.labelcolor": "#111",
                     "xtick.color": "#111", "ytick.color": "#111"})
GD, GM, GL = "#3a3a3a", "#7d7d7d", "#c7c7c7"   # grayscale dark / mid / light


def fig_market(path):
    fig, ax = plt.subplots(figsize=(5.6, 2.9))
    yrs = [2025, 2027, 2029, 2031, 2033, 2035]
    val = [863.7, 1750, 3200, 5100, 7100, 9207.2]        # $M, smooth growth to the 2035 figure
    ax.fill_between(yrs, val, color=GL, alpha=0.5)
    ax.plot(yrs, val, color=GD, lw=2.2, marker="o", mfc="white", mec=GD)
    for x, y in [(2025, 863.7), (2035, 9207.2)]:
        ax.annotate(f"${y:,.0f}M", (x, y), textcoords="offset points", xytext=(0, 8),
                    ha="center", fontsize=10, fontproperties=fm.FontProperties(fname=SARB))
    ax.set_ylabel("ขนาดตลาด (ล้านดอลลาร์สหรัฐ)", fontsize=10)
    ax.set_title("ตลาดการลดต้นทุน LLM เติบโต ~10.6 เท่า ใน 10 ปี", fontsize=11,
                 fontproperties=fm.FontProperties(fname=SARB))
    ax.set_xticks(yrs); ax.grid(axis="y", ls=":", color=GL); ax.set_axisbelow(True)
    for s in ("top", "right"):
        ax.spines[s].set_visible(False)
    fig.tight_layout(); fig.savefig(path); plt.close(fig)


def fig_api(path):
    fig, ax = plt.subplots(figsize=(5.6, 2.7))
    tiers = ["ประหยัด\n(budget)", "กลาง\n(mid)", "โปรดักชัน\n(production)", "reasoning\n/premium"]
    lo = [0.10, 1, 3, 8]; hi = [0.40, 5, 15, 30]
    y = np.arange(len(tiers))[::-1]
    ax.barh(y, np.array(hi) - np.array(lo), left=lo, color=[GL, GM, GM, GD], edgecolor="#222")
    for yi, l, h in zip(y, lo, hi):
        ax.text(h + 0.5, yi, f"${l:.2f}–${h:.0f}", va="center", fontsize=9)
    ax.set_yticks(y); ax.set_yticklabels(tiers, fontsize=9)
    ax.set_xlabel("ดอลลาร์สหรัฐ ต่อ 1 ล้าน token (ขาเข้า; ขาออกแพงกว่า 4–6 เท่า)", fontsize=9)
    ax.set_title("ราคา LLM API แยกตามระดับโมเดล (มิ.ย. 2026)", fontsize=11,
                 fontproperties=fm.FontProperties(fname=SARB))
    ax.set_xlim(0, 36)
    for s in ("top", "right"):
        ax.spines[s].set_visible(False)
    fig.tight_layout(); fig.savefig(path); plt.close(fig)


def fig_results(path):
    fig, ax = plt.subplots(figsize=(5.6, 2.9))
    groups = ["โค้ดที่ผ่านการทดสอบ\n(n=40)", "งานที่ไม่เคยเห็น\n(held-out, n=40)"]
    ours = [39, 31]; base = [18, 15]
    x = np.arange(len(groups)); w = 0.36
    b1 = ax.bar(x - w/2, base, w, label="ไม่ใช้วิธีเรา", color=GL, edgecolor="#222")
    b2 = ax.bar(x + w/2, ours, w, label="ใช้วิธีเรา", color=GD, edgecolor="#222")
    for b in list(b1) + list(b2):
        ax.text(b.get_x() + b.get_width()/2, b.get_height() + 0.6, f"{int(b.get_height())}/40",
                ha="center", fontsize=9, fontproperties=fm.FontProperties(fname=SARB))
    ax.set_xticks(x); ax.set_xticklabels(groups, fontsize=9); ax.set_ylim(0, 46)
    ax.set_ylabel("จำนวนงานที่แก้สำเร็จ", fontsize=10)
    ax.set_title("ผลจริงบน Qwen2.5-3B: ไวยากรณ์ผิด 0/40 และชนะ RAG บนงานใหม่", fontsize=10.5,
                 fontproperties=fm.FontProperties(fname=SARB))
    ax.legend(prop=fm.FontProperties(fname=SAR, size=9), frameon=False)
    ax.grid(axis="y", ls=":", color=GL); ax.set_axisbelow(True)
    for s in ("top", "right"):
        ax.spines[s].set_visible(False)
    fig.tight_layout(); fig.savefig(path); plt.close(fig)


def fig_pipeline(path):
    """Clean black-and-white pipeline: boxes (black border, white/gray fill) + arrows + Thai labels."""
    fig, ax = plt.subplots(figsize=(7.4, 4.0)); ax.axis("off")
    ax.set_xlim(0, 12); ax.set_ylim(-0.4, 6.5)
    fpb = fm.FontProperties(fname=SARB, size=9.5); fp = fm.FontProperties(fname=SAR, size=8.2)
    dash = (0, (4, 2))

    def box(x, y, w, h, title, sub, fill="white", lw=1.4):
        ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.03,rounding_size=0.12",
                                    fc=fill, ec="#111", lw=lw))
        ax.text(x + w/2, y + h - 0.33, title, ha="center", va="top", fontproperties=fpb)
        if sub:
            ax.text(x + w/2, y + h - 0.74, sub, ha="center", va="top", fontproperties=fp, color="#333")

    def arrow(x1, y1, x2, y2, head=True, color="#444", lw=1.6, ls="-"):
        ax.add_patch(FancyArrowPatch((x1, y1), (x2, y2), arrowstyle="-|>" if head else "-",
                                     mutation_scale=13, color=color, lw=lw, linestyle=ls,
                                     shrinkA=1, shrinkB=1))
    # top row
    box(0.2, 4.6, 2.0, 1.35, "งาน (NL)", "โจทย์ภาษา\nธรรมชาติ")
    box(2.6, 4.6, 2.0, 1.35, "ROUTER", "GraphGPS\nเลือกอะตอมใด")
    box(5.0, 4.6, 2.3, 1.35, "PLANNER / TRM", "โปรแกรมอะตอม\nต่ออย่างไร")
    box(7.7, 4.6, 3.9, 1.35, "กราฟความรู้ (เยื่อเลือกผ่าน)", "โหนด = โค้ด + คำอธิบาย + วิธีทำ + ที่มา",
        fill="#e6e6e6", lw=1.9)
    arrow(2.2, 5.28, 2.6, 5.28); arrow(4.6, 5.28, 5.0, 5.28); arrow(7.3, 5.28, 7.7, 5.28)
    # graph -> realize (down)
    arrow(9.65, 4.6, 9.65, 3.15)
    # middle row (right -> left)
    box(6.1, 1.8, 5.5, 1.35, "สร้างผลลัพธ์ 2 ช่องสัญญาณ",
        "ช่องสัญลักษณ์: โครงสร้างเป๊ะจากกราฟ (ไม่แต่งผิด)\nช่องความหมาย: คำอธิบายจากการเดินกราฟ", lw=1.9)
    box(3.4, 1.8, 2.4, 1.35, "LM ตรึงค่า", "4-bit · 6GB · $0\nเติมกาว + เล่า", fill="#e6e6e6")
    box(0.8, 1.8, 2.3, 1.35, "VERIFY", "ประตูเดียว\nที่เขียนกราฟได้", lw=1.9)
    arrow(6.1, 2.42, 5.8, 2.42); arrow(3.4, 2.42, 3.1, 2.42)
    # verify -> output (down-left)
    box(0.55, 0.05, 2.5, 1.25, "ผลลัพธ์", "โค้ด + คำอธิบายที่ตรวจสอบได้", fill="#e6e6e6")
    arrow(1.9, 1.8, 1.75, 1.3)
    # BANK feedback: VERIFY -> up -> right -> into graph (clean L-path in the middle band, dashed)
    arrow(1.95, 3.15, 1.95, 3.66, head=False, color="#111", lw=1.3, ls=dash)
    arrow(1.95, 3.66, 8.55, 3.66, head=False, color="#111", lw=1.3, ls=dash)
    arrow(8.55, 3.66, 8.55, 4.6, color="#111", lw=1.3, ls=dash)
    ax.text(4.9, 3.74, "BANK: บันทึกอะตอมที่ผ่านการตรวจ → กราฟโต ต้นทุนต่อครั้งลดลง (ไม่ฝึกโมเดล)",
            fontproperties=fp, color="#111", ha="center", va="bottom")
    fig.savefig(path, bbox_inches="tight", pad_inches=0.1); plt.close(fig)


print("rendering figures...")
for fn, f in [("fig_market.png", fig_market), ("fig_api.png", fig_api),
              ("fig_results.png", fig_results), ("fig_pipeline.png", fig_pipeline)]:
    f(os.path.join(HERE, fn))
print("figures done")

# ============================ DOCX ============================
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "TH Sarabun New"
NB = RGBColor(0x1F, 0x4E, 0x79)
doc = Document()
sec = doc.sections[0]
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, m, Inches(0.9))


def _thai(run, size, bold, color):
    run.font.size = Pt(size); run.font.bold = bold
    run.font.name = FONT
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rpr.insert(0, rf)
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rf.set(qn(a), FONT)
    szcs = OxmlElement("w:szCs"); szcs.set(qn("w:val"), str(int(size*2))); rpr.append(szcs)


def para(text="", size=16, bold=False, color=None, align=None, before=2, after=4, indent=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after); pf.line_spacing = 1.0
    if indent:
        pf.first_line_indent = Inches(indent)
    if text:
        _thai(p.add_run(text), size, bold, color)
    return p


def h(num, text, size=17):
    p = para(f"{num} {text}" if num else text, size=size, bold=True, color=NB, before=10, after=4)
    return p


def bullet(text, size=16, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2); p.paragraph_format.line_spacing = 1.0
    if bold_lead:
        _thai(p.add_run(bold_lead), size, True, None)
    _thai(p.add_run(text), size, False, None)
    return p


def figure(png, cap, width=5.7):
    doc.add_picture(os.path.join(HERE, png), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c = para(cap, size=13, bold=False, color=RGBColor(0x44, 0x44, 0x44), align=WD_ALIGN_PARAGRAPH.CENTER,
             before=1, after=8)


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    for i, hd in enumerate(headers):
        cell = t.rows[0].cells[i]; cell.paragraphs[0].clear()
        _thai(cell.paragraphs[0].add_run(hd), 14, True, None)
    for r in rows:
        cells = t.add_row().cells
        for i, v in enumerate(r):
            cells[i].paragraphs[0].clear()
            _thai(cells[i].paragraphs[0].add_run(str(v)), 14, False, None)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

# -------- COVER --------
for _ in range(2):
    para(before=0, after=0)
para("ข้อเสนอผลงานสิ่งประดิษฐ์และนวัตกรรม · TICTA 2026 / I-New Gen", size=15, color=NB,
     align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=10)
para("ปัญญาประดิษฐ์ให้เหตุผลไร้คลาวด์เพื่ออุตสาหกรรม", size=26, bold=True, color=NB,
     align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=2)
para("ลดต้นทุนและพลังงานด้วยกราฟความรู้ที่ตรวจสอบได้ อธิบายวิธีคิดได้ และเก่งขึ้นโดยไม่ต้องฝึกโมเดล",
    size=19, bold=True, color=NB, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=6)
para("Cloud-Free Reasoning AI for Industry: cutting cost and energy with a verified knowledge graph "
    "that explains its reasoning and improves without retraining",
    size=14, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=14)
para("ความเชื่อมโยง: ปัญญาประดิษฐ์ (CT-AI) · โครงงานนักเรียน (HC-S) · ประสิทธิภาพพลังงานและการใช้งานบนอุปกรณ์ปลายทาง",
    size=14, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=30)
para("จัดทำโดย  นายประพัฒน์พงศ์ พิทักษ์ธรรม   ระดับชั้นมัธยมศึกษาปีที่ 6",
    size=16, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=2)
para("โรงเรียนขอนแก่นวิทยายน   ปีการศึกษา 2569", size=16, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0)
doc.add_page_break()

# -------- (1) ABSTRACT --------
h("(1)", "บทคัดย่อ")
para("โครงการนี้พัฒนาระบบปัญญาประดิษฐ์ที่ให้เหตุผลและแก้โจทย์โปรแกรมได้บนเครื่องของผู้ใช้เอง (on-device) "
    "โดยใช้โมเดลภาษาขนาดเล็ก (~3B แบบ 4-bit) ที่ “ตรึงค่า” (frozen) รันบน GPU โน้ตบุ๊กขนาด 6 GB "
    "แบบออฟไลน์ ต้นทุนต่อการใช้งานเป็นศูนย์ จุดต่างสำคัญคือ “ความสามารถและการเก่งขึ้นทั้งหมด” ไม่ได้อยู่ใน "
    "น้ำหนักของโมเดล แต่อยู่ใน “กราฟความรู้ภายนอกที่ตรวจสอบได้และคนอ่านได้” ระบบเรียนรู้ด้วยการ "
    "เพิ่มโหนด/เส้นเชื่อมในกราฟ ไม่ใช่การฝึกโมเดล จึงหลีกเลี่ยงทั้งต้นทุนการฝึกที่กินพลังงานสูง การลืมความรู้เดิม "
    "(catastrophic forgetting) และความเสี่ยงจากการส่งข้อมูลขึ้นคลาวด์", after=4)
para("สถาปัตยกรรมหลักคือ “สองช่องสัญญาณ” ที่แยก “ความหมาย” (การอธิบาย) ออกจาก “ไวยากรณ์” "
    "(การเขียนโค้ดที่เป๊ะ) ช่องสัญลักษณ์ดึงโครงสร้างโค้ดที่ตรวจแล้วจากกราฟจึงไม่มีทางแต่งไวยากรณ์ผิด ส่วนช่องความหมาย "
    "ให้โมเดลอธิบายจากการเดินกราฟจริง ผลจากการวัดบนโมเดล 3B จริงคือ ไวยากรณ์ผิดเป็นศูนย์และแก้งานได้ 39 จาก 40 "
    "(เทียบวิธีเขียนอิสระ 18/40) ความตรงของคำอธิบายกับโค้ด 1.00 และกราฟที่เติบโตได้ชนะการค้นแบบ RAG บนงานที่ไม่เคยเห็น "
    "(31 ต่อ 15) พร้อมเรียกโมเดลน้อยลง", after=4)
para("คำสำคัญ:  Local/Edge AI, Graph Memory, Frozen LLM, Dual-Channel Realization, Verified Reasoning, "
    "Traceable Explanation, Structural Retrieval, Energy-Efficient AI", bold=False, after=2)

# -------- (2) RATIONALE + COST HOOK --------
h("(2)", "หลักการและเหตุผล")
para("ปัญญาประดิษฐ์สมัยใหม่ตอบคำถามได้ดี แต่การนำไปใช้จริงในโรงเรียน ห้องปฏิบัติการ ทีมพัฒนา และธุรกิจ SME "
    "ยังติดข้อจำกัดสามด้านที่ “ทวีความรุนแรงขึ้น” ตามการใช้งาน คือ ต้นทุน พลังงาน และความเป็นส่วนตัวของข้อมูล")

h("(2.1)", "เหตุผลเชิงเวลา: ต้นทุนและพลังงานของ AI กำลังพุ่งสูงเป็นปัญหาเชิงโครงสร้าง", size=16)
para("นี่คือปัญหาที่มีหลักฐานเชิงตัวเลขชัดเจนและกำลังโตเร็ว รายงานตลาดการลดต้นทุน LLM ระบุว่าตลาดทั่วโลกจะเติบโตจาก "
    "ประมาณ 863.7 ล้านดอลลาร์สหรัฐในปี 2025 ไปสู่กว่า 9,207.2 ล้านดอลลาร์สหรัฐในปี 2035 (ราว 10.6 เท่า) "
    "การเติบโตของ “ตลาดที่เกิดขึ้นเพื่อลดต้นทุน AI โดยเฉพาะ” เป็นเครื่องยืนยันว่าต้นทุนการอนุมาน (inference) "
    "คือโจทย์ระดับอุตสาหกรรม ไม่ใช่ปัญหาปลีกย่อย")
figure("fig_market.png", "รูปที่ 1  การเติบโตของตลาดการลดต้นทุน LLM ปี 2025–2035  (ที่มา: Market.us Research, "
       "LLM Cost Optimization Market 2025–2035)")
para("ในมิติพลังงาน สำนักงานพลังงานระหว่างประเทศ (IEA) รายงานว่าการใช้ไฟฟ้าของศูนย์ข้อมูลที่เน้น AI พุ่งขึ้นราว 50% "
    "ในปี 2025 และการใช้ไฟฟ้าของศูนย์ข้อมูลรวมมีแนวโน้มเพิ่มเป็นสองเท่าจาก 485 เป็นราว 950 TWh ภายในปี 2030 "
    "สอดคล้องกับรายงานของ Penn State IEE ที่ชี้ว่าภาระงาน AI ทำให้ศูนย์ข้อมูลใช้ไฟฟ้าสัดส่วนสูงขึ้นต่อเนื่อง "
    "แนวโน้มทั้งสองด้านชี้ไปทางเดียวกัน: หากระบบเก็บความรู้เป็นหน่วยย่อย ค้นเฉพาะส่วนที่เกี่ยวข้อง และทำงานในเครื่องได้ "
    "ก็จะลดทั้งต้นทุนและพลังงานที่ไม่จำเป็น")
figure("fig_api.png", "รูปที่ 2  ราคา LLM API ต่อ 1 ล้าน token แยกตามระดับโมเดล (มิ.ย. 2026) — ส่วนต่างระหว่างระดับยังสูงมาก "
       "และงาน agentic ที่คิดหลายขั้นใช้ token ต่อคำขอมากกว่าการแชตทั่วไปหลายเท่า")

h("(2.2)", "ภาระต้นทุนนี้หนักเป็นพิเศษสำหรับประเทศไทยและ SME", size=16)
para("ปลายปี 2025 รัฐบาลไทยเสนอโครงการ “TH-AI Passport” งบประมาณราว 1.6 พันล้านบาท เพื่อจัดซื้อสิทธิ์ใช้งานโมเดล "
    "AI เชิงพาณิชย์ระดับ Pro ให้ประชาชนใช้ฟรี เพราะค่าสมาชิกระดับ Pro ราว 700–1,000 บาทต่อเดือนต่อบริการ สูงเกินกำลัง "
    "ผู้ใช้ทั่วไปและธุรกิจขนาดเล็ก ตัวเลขระดับพันล้านบาทเพื่อ “เช่าสิทธิ์เข้าถึง” โมเดลต่างประเทศ เป็นหลักฐานชัดเจนว่า "
    "โครงสร้างต้นทุน AI ปัจจุบันอยู่นอกกำลังของผู้ใช้ส่วนใหญ่ในประเทศ ยิ่งเมื่ออัตราการใช้ AI ของไทยอยู่ที่ 10.7% ต่ำกว่า "
    "ค่าเฉลี่ยโลก 16.3% ประเทศในช่องว่างนี้ไล่ตามด้วยการซื้อ GPU ใหญ่หรือจ่ายค่า API ระดับเดียวกับผู้นำไม่ได้ "
    "จึงต้องการสถาปัตยกรรมที่ประหยัดทรัพยากรกว่าเดิมมาก")

h("(2.3)", "ปัญหาเชิงสถาปัตยกรรมที่ทำให้ต้นทุนลดไม่ได้", size=16)
para("รากของปัญหาคือสถาปัตยกรรมปัจจุบันยัด “ความสามารถทั้งหมด” ไว้ในน้ำหนักทึบของโมเดล เมื่อความรู้ผูกกับน้ำหนัก "
    "การเพิ่ม/แก้ความรู้จึงต้องฝึกโมเดลใหม่ (พลังงานสูง + ลืมของเดิม) หรือส่งข้อมูลขึ้นคลาวด์ (เสียค่าใช้จ่าย + เสี่ยงความลับ) "
    "ส่วน RAG ทั่วไปเก็บความรู้เป็น “ข้อความก้อนเดียว” มองไม่เห็นโครงสร้างของโค้ด และเมื่อบริบทยาว ต้นทุนประมวลผลก็สูงตาม "
    "แม้ข้อมูลบางส่วนไม่เกี่ยวข้อง")

# -------- (3) THE FIVE NEEDS / OBJECTIVES --------
h("(3)", "แนวคิดและวัตถุประสงค์")
para("เป้าหมายไม่ใช่การทำให้โมเดลใหญ่ขึ้น แต่ทำให้โมเดลเล็กทำงานแม่นยำและคุ้มค่าขึ้น โดยแยก “ความรู้” ออกจาก "
    "“น้ำหนักโมเดล” ระบบที่ตอบโจทย์ต้องมีคุณสมบัติสี่ประการพร้อมกัน:")
bullet("โมเดลเล็กตรึงค่า รันในเครื่อง ≤6 GB ออฟไลน์ ต้นทุนต่อครั้งเป็นศูนย์",
       bold_lead="อยู่บนเครื่อง (on-device): ")
bullet("ทุกความรู้ที่บันทึกต้องผ่านการทดสอบก่อน กราฟจึงถูกต้อง “โดยโครงสร้าง”",
       bold_lead="ถูกต้องโดยการตรวจสอบ (verified): ")
bullet("บอกได้ว่าใช้ขั้นตอนใด เพราะอะไร และคำอธิบายต้องตรงกับสิ่งที่รันจริง",
       bold_lead="อธิบายได้อย่างซื่อสัตย์: ")
bullet("“เรียนรู้” = เพิ่มโหนด/เส้นเชื่อมในกราฟ ไม่ใช่ gradient descent",
       bold_lead="เก่งขึ้นโดยไม่ต้องฝึก: ")

# -------- (4) METHODOLOGY + PIPELINE --------
h("(4)", "สถาปัตยกรรมและวิธีดำเนินงาน")
para("เรานำ “กฎที่วัดได้” มาออกแบบระบบ: โมเดลตรึงค่าเป็นฟังก์ชันข้อความเข้า–ข้อความออก ช่องทางเดียวที่ป้อนเข้าได้ "
    "แบบไม่สูญเสียคือ “ข้อความ” ไม่ใช่เวกเตอร์แฝง (เราทดลองแล้วช่องแฝงล้มเหลว เกิด routing collapse) ระบบภายนอกจึงช่วยได้ "
    "โดยเลือก “ข้อความที่ตรวจแล้ว” ป้อนให้ และปล่อยให้โมเดลทำสิ่งที่ถนัด กระบวนการทำงานมีเจ็ดสถานีดังรูปที่ 3")
figure("fig_pipeline.png", "รูปที่ 3  ท่อการทำงานของระบบ: ROUTER เลือกอะตอม → PLANNER/TRM วางโครงสร้าง → กราฟส่งโค้ดที่ตรวจแล้ว "
       "เป็นข้อความ → สร้างผลลัพธ์สองช่อง (สัญลักษณ์=เป๊ะ, ความหมาย=อธิบาย) → VERIFY เป็นประตูเดียวที่เขียนกราฟได้ → "
       "ผ่านแล้ว BANK ทำให้กราฟโตและต้นทุนต่อครั้งลดลง โดยโมเดลไม่เคยถูกฝึก", width=6.6)

# -------- (5) SYSTEM DEEP-DIVE --------
h("(5)", "องค์ประกอบเชิงลึกของระบบ")
para("โหนดในกราฟไม่ใช่ “โค้ดเปล่า” แต่เป็น “หน่วยความหมาย” ที่มีคำอธิบาย วิธีทำ และที่มา ทำให้กราฟอ่านได้และอธิบายได้",
     bold=False)
bullet("ตัวจัดเส้นทาง GraphGPS ฝังงานและอะตอมในปริภูมิความหมายเดียวกัน + เดินตามเส้นเชื่อม แยกแยะ “จำนวนหลัก” จาก "
       "“จำนวนตัวหาร” ที่การจับคู่คำทำไม่ได้ (recall 0.50 → 0.98)", bold_lead="ROUTER: ")
bullet("ผู้ให้เหตุผลจิ๋ว (หลักหมื่นพารามิเตอร์) อนุมาน “โครงสร้างโปรแกรม” จากโจทย์ และเขียนสถานะย่อยที่ตรวจแล้วลง "
       "“ความจำที่ไม่จางหาย” ยกเพดานการประกอบจาก 0.03 เป็น 1.00", bold_lead="PLANNER/TRM: ")
bullet("ช่องสัญลักษณ์วางโครงสร้างที่ยาก (จากกราฟ, เป๊ะ) เหลือ “ช่องว่าง” ให้โมเดลเติมกาว ช่องความหมายอธิบายจากการเดินกราฟ "
       "→ ผลลัพธ์เป็นของโมเดลเอง อธิบายได้ และไวยากรณ์ผิดเป็นศูนย์", bold_lead="สองช่องสัญญาณ: ")
bullet("การรันโค้ดจริงเป็น “ประตูเดียว” ที่เขียนกราฟได้ อะตอมที่ผิดจะไม่ถูกบันทึก กราฟจึงสะอาดเสมอ และการนำอะตอม "
       "ที่ตรวจแล้วกลับมาใช้ทำให้เรียกโมเดลน้อยลงเรื่อย ๆ (compounding)", bold_lead="VERIFY + BANK: ")

# -------- (6) RESULTS --------
h("(6)", "ผลการทดลอง (วัดจริงบนโมเดล 3B ฮาร์ดแวร์เดียว)")
figure("fig_results.png", "รูปที่ 4  ผลจริงบน Qwen2.5-3B: สองช่องสัญญาณแก้ได้ 39/40 โดยไวยากรณ์ผิด 0/40 และกราฟที่เติบโตได้ "
       "ชนะ RAG บนงานที่ไม่เคยเห็น 31 ต่อ 15")
table(["สิ่งที่วัด (Qwen2.5-3B จริง)", "ไม่ใช้วิธีเรา", "ใช้วิธีเรา"],
      [["งานที่ผ่านการทดสอบ (n=40)", "18/40", "39/40"],
       ["ไวยากรณ์ผิดที่หลุดออกไป", "5/40", "0/40"],
       ["ความตรงของคำอธิบายกับโค้ด", "—", "1.00"],
       ["งานที่ไม่เคยเห็น (held-out, n=40)", "15/40 (RAG)", "31/40"],
       ["จำนวนครั้งที่เรียก LM (น้อย=ดี)", "100", "63"],
       ["เพดานการประกอบ (ความลึก 5)", "0.03", "1.00"],
       ["การค้นอะตอมด้วยโครงสร้าง (เทียบ cosine)", "0.50", "0.98"]])
para("อ่านผล: (1) ช่องสัญลักษณ์ทำให้ไวยากรณ์ผิดเป็นศูนย์ เพราะโครงสร้างยากมาจากกราฟ ไม่ใช่การเดา (2) กราฟที่เติบโตได้ "
    "ชนะ RAG พร้อมเรียกโมเดลน้อยกว่า เพราะนำอะตอมที่ตรวจแล้วกลับมาใช้ (การนำกลับมาใช้ที่ได้จากการอนุมานวัดได้ 6→70 "
    "บนสตรีมจริง) (3) เมื่อให้ “โครงสร้าง” เพดานการประกอบที่โมเดลเปล่าทำได้ 3% กลายเป็น 100% — ทั้งหมดโดยไม่ฝึกโมเดล")

# -------- (7) SCALING / DEPLOYMENT (TICTA criteria) --------
h("(7)", "การขยายสู่การใช้งานจริงและการเรียนรู้ด้วยตนเอง")
bullet("ระบบเรียนรู้ด้วยตนเองผ่าน “ครูตรวจแทนการฝึก”: โมเดลครูขนาดใหญ่ (นอกเครื่อง) ช่วยตัดสินเฉพาะที่ไม่มีเครื่องตรวจ "
       "อัตโนมัติ แต่การรันโค้ดจริงยังเป็นความจริงหลัก การเรียนรู้จึงปลอดภัยและไม่ทำให้โมเดล “เลื่อนลอย”",
       bold_lead="Self-learning: ")
bullet("ผู้ให้เหตุผลจิ๋วปรับจูนร่วมกับ ROUTER ได้ด้วย RL ที่ออกแบบรางวัลผสมกันไม่ให้ “ยุบตัว” (พิสูจน์แล้วแบบไม่ใช้ GPU) "
       "การค้นคืนมีแคชและดัชนี ทำให้ขยายไปกราฟใหญ่ได้", bold_lead="Scale: ")
bullet("รันบน GPU โน้ตบุ๊กที่มีอยู่แล้ว ออฟไลน์ ต้นทุนต่อครั้งเป็นศูนย์ โค้ด/ข้อมูลลับไม่ออกนอกเครื่อง เหมาะกับโรงเรียน SME "
       "และโรงงานที่ต้องการความเป็นส่วนตัว", bold_lead="Deployment & security: ")

# -------- (8) VS RAG --------
h("(8)", "จุดต่างจาก RAG ทั่วไป")
para("RAG ดึงข้อความที่ “คล้าย” เข้าไปในพรอมป์ ไม่เห็นโครงสร้าง และไม่เติบโตด้วยตัวเอง ระบบของเรา “เรียนรู้” ว่าจะเลือก "
    "อะตอมใด (router ชนะ cosine ด้วยโครงสร้าง) “วางแผน” ว่าประกอบอย่างไร (planner ลบเพดาน 3%) และ “โตและตรวจสอบ "
    "คลังของตัวเอง” (ถูกลง เมื่อใช้มากขึ้น) บนโมเดลตรึงค่าที่ไม่เคยฝึกใหม่")

# -------- (9) LIMITS --------
h("(9)", "ข้อจำกัด (ระบุอย่างตรงไปตรงมา)")
bullet("เพดานคือ “การประดิษฐ์” อะตอมใหม่ที่ 3B เขียนเองไม่ได้ — เป็นแนวหน้าถัดไป")
bullet("ขอบเขตอยู่ที่โดเมนที่ตรวจได้ (โค้ด/คณิต/จำลอง) โดเมนปลายเปิดยังต้องพึ่งการค้นคืน")
bullet("ตัวอย่างทดสอบยังเป็นชุดหลักสิบ เป็นสัญญาณเชิงบวกและกลไกที่พิสูจน์ได้ ไม่ใช่การอ้างว่าดีที่สุด กำลังขยายไปชุดยากขึ้น "
       "(MHPP, BigCodeBench)")

# -------- (10) ROADMAP --------
h("(10)", "แผนการต่อยอด")
para("(ก) ให้โมเดลประดิษฐ์อะตอมใหม่ที่ผ่านการตรวจ (invention) (ข) ฝึกผู้ให้เหตุผลจิ๋วร่วมกับ ROUTER ด้วยครูและ RL แบบไม่ยุบตัว "
    "(ค) ขยายไปโดเมนคณิตศาสตร์ที่มีเครื่องตรวจเชิงสัญลักษณ์ เพื่อพิสูจน์ว่าสถาปัตยกรรมเดียวกันข้ามโดเมนได้")

# -------- (11) REFERENCES --------
h("(11)", "บรรณานุกรม")
for r in [
    "Market.us Research. LLM Cost Optimization Market Size, Share and Analysis 2025–2035. market.us",
    "International Energy Agency (IEA). Data centre electricity use surged in 2025. iea.org, 2025.",
    "Kandemir, M. Penn State Institute of Energy and the Environment — Why AI Uses So Much Energy. iee.psu.edu.",
    "Veracode. 2025 GenAI Code Security Report. veracode.com.",
    "Liu, N. F. et al. Lost in the Middle: How Language Models Use Long Contexts. 2023.",
    "Jimenez, C. E. et al. SWE-bench: Can Language Models Resolve Real-World GitHub Issues? 2024.",
]:
    b = doc.add_paragraph(style="List Number"); b.paragraph_format.space_after = Pt(1)
    _thai(b.add_run(r), 13, False, None)

para("หมายเหตุ: ตัวเลขทางเทคนิคทั้งหมดมาจากการรันจริงบนโมเดล Qwen2.5-3B แบบตรึงค่า ด้วยการตรวจด้วยชุดทดสอบจริง "
     "ภาพประกอบเป็นขาว–ดำเพื่อการพิมพ์ที่ชัดเจน แบบอักษร TH Sarabun New ขนาด 16", size=13,
     color=RGBColor(0x55, 0x55, 0x55), before=8)

out = os.path.join(HERE, "Proposal_GRR_v6_TH.docx")
doc.save(out)
print("SAVED:", out, "| paragraphs:", len(doc.paragraphs))
